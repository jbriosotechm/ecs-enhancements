"""
Requirement :
1.Convert Excel Report to Word
2. Generate Test Case by Test Case report
3. Generate entire report also
4. Write only specific columns

Logic:
1. Iterate over excel
2. Go on creating tables and writing data at the same time

"""
import sys,os
import string
import userConfig
import systemConfig
from docx import Document
from docx.shared import Inches,Pt,Length
import openpyxl
from openpyxl import styles
from openpyxl.utils import get_column_letter,column_index_from_string
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import time
import os,sys
import datetime
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import comtypes.client
import re,traceback
from Configuration_Browser import *
import Configuration_Browser
import random,fileinput

def mainWithMultipleArgs(arg1,arg2,arg3,arg4,arg5,arg6):
    #to support legacy
    main(arg2,arg6)




def getCellValueUsingColumnNumber(row,colName):
    try:

        val=systemConfig.sheet_obj.cell(row=row, column=colName).value
        if val is None or val.strip() in ["","None"]:
            return None
    except:
        return None


def getCellValueUsingCoordinate(coordinate):
    try:

        val=systemConfig.sheet_obj[coordinate].value
        if val is None or val.strip() in ["","None"]:
            return None
    except:
        return None

    return val


def getCellValue(row,colName):
    val=None
    try:
        val=systemConfig.sheet_obj["{0}{1}".format(colName,row)].value
        if val is None or val.strip() in ["","None"]:
            return None
    except:
        traceback.print_exc()
        #return None

    return val


def testCaseColumnIsEmpty(rowNumber):
    if getCellValue(rowNumber,systemConfig.column_testCaseId) is not None:
        return False

    return True

def testStepColumnIsEmpty(rowNumber):
    if getCellValue(rowNumber,systemConfig.column_testStepNumber) is not None:
        return False

    return True

def getMaxRow():
    #return systemConfig.sheet_obj.max_row

    #logic for another day if base function doesnt work
    lastLimit=systemConfig.sheet_obj.max_row
    startLimit=2
    while startLimit<lastLimit:

        if testCaseColumnIsEmpty(startLimit) and testStepColumnIsEmpty(startLimit):
            break
        startLimit+=1
    print("Max rows : {0}".format(startLimit-1))
    #time.sleep(5)
    return startLimit



def printRange(range):
    for row in range:
        for cell in row:
            print "value at {0}: {1}".format(cell.coordinate,cell.value)

        #new row
        print "###########"

def getExcelRangeWithoutHeaders():
    wb_obj = openpyxl.load_workbook(userConfig.excelFileName)
    systemConfig.sheet_obj = wb_obj.active
    maxRows=getMaxRow()
    range=systemConfig.sheet_obj["A2":"{0}{1}".format(systemConfig.LastColumn,maxRows)]
    return range

def rowIsTestCase(rowObject):
    for cell in rowObject:
        if str(systemConfig.column_testCaseId) in cell.coordinate:
            if getCellValueUsingCoordinate(cell.coordinate) is not None:
                return True
            else:
                return False


def getCellValueFromRowObject(rowObject,columnNameToFind):
    try:
        #print "RowObject : ",rowObject
        for cell in rowObject:

            if columnNameToFind in cell.coordinate:
                return getCellValueUsingCoordinate(cell.coordinate)
    except:
        traceback.print_exc()
    return None

def ColorTheRow(table,hexcolorcode):
    cellNumber=0
    while cellNumber<5:
        data=r'<w:shd {} w:fill="'+hexcolorcode+r'"/>'
        table.rows[0].cells[cellNumber]._tc.get_or_add_tcPr().append(parse_xml(data.format(nsdecls('w'))))
        cellNumber+=1


def createHeaderForTestCase():
        systemConfig.document.add_paragraph('\n')
        table1 = systemConfig.document.add_table(rows=1, cols=5)
        table1.style = 'TableGrid'
        table1.allow_autofit = True
        headerCells = table1.rows[0].cells
        headerCells[0].paragraphs[0].add_run(systemConfig.name_testCaseId).bold = True
        headerCells[1].paragraphs[0].add_run(systemConfig.name_testDesc).bold = True
        headerCells[2].paragraphs[0].add_run(systemConfig.name_Status).bold = True
        headerCells[3].paragraphs[0].add_run(systemConfig.name_StartTime).bold = True
        headerCells[4].paragraphs[0].add_run(systemConfig.name_EndTime).bold = True
        ColorTheRow(table1, systemConfig.hexcolorcode_test_case_ref)



def createHeaderForTestStep():
        systemConfig.document.add_paragraph('\n')
        table1 = systemConfig.document.add_table(rows=1, cols=5)
        table1.style = 'TableGrid'
        table1.allow_autofit = True
        headerCells = table1.rows[0].cells
        headerCells[0].paragraphs[0].add_run(systemConfig.name_testStepNumber).bold = True
        headerCells[1].paragraphs[0].add_run(systemConfig.name_testStepDesc).bold = True
        headerCells[2].paragraphs[0].add_run(systemConfig.name_ExpectedResult).bold = True
        headerCells[3].paragraphs[0].add_run(systemConfig.name_ActualResult).bold = True
        headerCells[4].paragraphs[0].add_run(systemConfig.name_Status).bold = True
        ColorTheRow(table1, systemConfig.hexcolorcode_test_step_ref)



def populateTestStepDetails(rowObject):
        table1 = systemConfig.document.add_table(rows=1, cols=5)
        table1.style = 'TableGrid'
        table1.allow_autofit = True
        headerCells = table1.rows[0].cells

        #print "Row Object : ",rowObject
        testStepNumber=getCellValueFromRowObject(rowObject,systemConfig.column_testStepNumber)
        testDesc=getCellValueFromRowObject(rowObject,systemConfig.column_testDescription)
        expectedResult=getCellValueFromRowObject(rowObject,systemConfig.column_expectedResult)
        actualResult=getCellValueFromRowObject(rowObject,systemConfig.column_actualResult)
        status=getCellValueFromRowObject(rowObject,systemConfig.column_stepStatus)

        writeTestStep(testStepNumber,testDesc,expectedResult,actualResult,status)

        headerCells[0].paragraphs[0].add_run(testStepNumber).bold = False
        headerCells[1].paragraphs[0].add_run(testDesc).bold = False
        headerCells[2].paragraphs[0].add_run(expectedResult).bold = False
        headerCells[3].paragraphs[0].add_run(actualResult).bold = False
        headerCells[4].paragraphs[0].add_run(status).bold = False


def populateTestCaseDetails(rowObject):

        table1 = systemConfig.document.add_table(rows=1, cols=5)
        table1.style = 'TableGrid'
        table1.allow_autofit = True
        headerCells = table1.rows[0].cells

        tcId=getCellValueFromRowObject(rowObject,systemConfig.column_testCaseId)
        testDesc=getCellValueFromRowObject(rowObject,systemConfig.column_testDescription)
        status=getCellValueFromRowObject(rowObject,systemConfig.column_stepStatus)
        startTime=getCellValueFromRowObject(rowObject,systemConfig.column_startTime)
        endTime=getCellValueFromRowObject(rowObject,systemConfig.column_endTime)

        if systemConfig.testCaseWrittenOnce:
            endTestCase()
        writeTestCase(tcId,testDesc,status,startTime,endTime)

        systemConfig.testCaseId=tcId
        systemConfig.testStepDesc=testDesc
        systemConfig.testStatus=status


        if "fail" in str(status).lower():
            systemConfig.fail_ctr+=1
            ColorTheRow(table1,systemConfig.hexcolorcode_fail)

        if "pass" in str(status).lower():
            systemConfig.pass_ctr+=1
            ColorTheRow(table1,systemConfig.hexcolorcode_pass)


        headerCells[0].paragraphs[0].add_run(tcId).bold = True
        headerCells[1].paragraphs[0].add_run(testDesc).bold = True
        headerCells[2].paragraphs[0].add_run(status).bold = True
        headerCells[3].paragraphs[0].add_run(startTime).bold = True
        headerCells[4].paragraphs[0].add_run(endTime).bold = True





def createTestCaseHeaderAndPopulate(rowObject):
    createHeaderForTestCase()
    populateTestCaseDetails(rowObject)

    return

def createTestStepHeaderAndPopulate(rowObject):
    createHeaderForTestStep()
    populateTestStepDetails(rowObject)
    return


def saveWordDocument(testCaseId,status,testStepDesc,outputPdfPath):
    print "Saving word doc"
    subfolderTestCaseWise=outputPdfPath+"\\TestCaseWise"

    if not os.path.exists(subfolderTestCaseWise):
        os.makedirs(subfolderTestCaseWise)

    absFilePath="{0}\\{1}".format(subfolderTestCaseWise,testCaseId+"_"+status)
    print "FilePath : {0}".format(absFilePath+".docx")
    systemConfig.document.save(absFilePath+".docx")

    file=open("{0}".format(absFilePath+".txt"),"w")
    file.write(testStepDesc)
    file.close()

def saveEntireWordDocument(outputPdfPath, outputPdfName="Report"):
    print "Saving word doc"
    absFilePath=outputPdfPath.replace(".xlsx",".docx")
    print("absFilePath:",absFilePath)
    #print "FilePath : {0}".format(absFilePath+".docx")
    systemConfig.document.save(absFilePath)




def addScreenshots(rowObject):
    screenshot_path=getCellValueFromRowObject(rowObject,systemConfig.column_screenshot)
    document=systemConfig.document

    image_height_inches1=Configuration_Browser.image_height_inches
    image_width_inches1=Configuration_Browser.image_width_inches

    if Configuration_Browser.deviceIsMobile==True:
        image_height_inches1=Configuration_Browser.image_height_inches_mobile
        image_width_inches1=Configuration_Browser.image_width_inches_mobile


    if screenshot_path is not None:
        try:
            if auto_image_height:
                document.add_paragraph('\n')
                document.add_picture(screenshot_path, width=Inches(2.5))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_paragraph('\n')
            else:
                document.add_paragraph('\n')
                document.add_picture(screenshot_path, width=Inches(image_width_inches1),
                                        height=Inches(image_height_inches1))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_paragraph('\n')
        except Exception, e:
            traceback.print_exc()



def initializeWordReport():

    systemConfig.pass_ctr=0
    systemConfig.fail_ctr=0
    document = Document()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(globe_logo, Inches(2), Inches(1))
    run.add_tab()
    run.add_picture(techm_logo, Inches(2), Inches(1))

    heading_para = document.paragraphs[-1]
    paragraph = document.add_paragraph('\nTest Summary Report')

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # paragraph.add_run('Test Summary Report', style = 'Title')
    paragraph.style = document.styles['Title']

    stats_paragraph = document.add_paragraph('ReplaceContent')
    stats_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stats_paragraph.style = document.styles['Heading 1']

    section = document.sections[-1]
    sections = document.sections
    section.page_width = Inches(11)
    section.page_height = Inches(17)

    # margins will also be a config parameter
    for section in sections:
        section.top_margin = Inches(0.05)
        section.bottom_margin = Inches(0.2)
        section.left_margin = Inches(0.2)
        section.right_margin = Inches(0.2)
        section.orientation = WD_ORIENT.LANDSCAPE

    systemConfig.document=document
    systemConfig.stats_paragraph=stats_paragraph

def formatParagraphs():
        para_counter = 0
        document=systemConfig.document
        for para in document.paragraphs:
            if para_counter > 1:
                paragraph_format = para.paragraph_format
                paragraph_format.line_spacing = 0.3
            else:
                paragraph_format = para.paragraph_format
                paragraph_format.line_spacing = 0.7
            para_counter += 1


def endWordReporting():
    total_test_cases = systemConfig.pass_ctr + systemConfig.fail_ctr
    result_str = "Total Test Cases : {0}      Passed : {1}      Failed : {2}".format(total_test_cases, systemConfig.pass_ctr,
                                                                                     systemConfig.fail_ctr)
    # stats_paragraph.add_run(result_str,style='Subtitle')
    systemConfig.stats_paragraph.text = result_str
    paragraph_format = systemConfig.stats_paragraph.paragraph_format
    paragraph_format.line_spacing = 2
    return



def endWordReportingTestCaseWise(status):
    total_test_cases = systemConfig.pass_ctr + systemConfig.fail_ctr

    if "fail" in status.lower():
        result_str = "Total Test Cases : {0}      Passed : {1}      Failed : {2}".format("1","0","1")
    else:
        result_str = "Total Test Cases : {0}      Passed : {1}      Failed : {2}".format("1","1","0")

    #systemConfig.stats_paragraph.add_run("Duration : TBD",style='Subtitle')
    systemConfig.stats_paragraph.text = result_str
    paragraph_format = systemConfig.stats_paragraph.paragraph_format
    paragraph_format.line_spacing = 2
    return

def generateTestCaseWiseReport(outputPdfPath):

    dataRangeToConvert=getExcelRangeWithoutHeaders()
    initializeWordReport()
    loggedOnce=False
    loopCounter=0
    testCaseLogged=False

    for eachExcelRow in dataRangeToConvert:
        print("Excel Row is : {0}".format(eachExcelRow))
        if rowIsTestCase(eachExcelRow):
            if loggedOnce:
                formatParagraphs()
                endWordReportingTestCaseWise(systemConfig.testStatus)
                saveWordDocument(systemConfig.testCaseId,systemConfig.testStatus,systemConfig.testStepDesc,outputPdfPath)
                initializeWordReport()
            createTestCaseHeaderAndPopulate(eachExcelRow)
            testCaseLogged=True


        else:
            loggedOnce=True
            if testCaseLogged is True:
                testCaseLogged=False
                createTestStepHeaderAndPopulate(eachExcelRow)
            else:
                populateTestStepDetails(eachExcelRow)
            addScreenshots(eachExcelRow)

    endTestCase()
    formatParagraphs()
    endWordReportingTestCaseWise(systemConfig.testStatus)
    saveWordDocument(systemConfig.testCaseId,systemConfig.testStatus,systemConfig.testStepDesc,outputPdfPath)



def formatUserInputs(output_pdf_path):
    #arg1,arg4,arg5 are redundant

    output_pdf_name="Report"


    if output_pdf_path[-1] == "\\" and output_pdf_path[-2] == "\\":
        output_pdf_path = output_pdf_path[:-1]
    print "output_pdf_path : ", output_pdf_path
    return (output_pdf_path,output_pdf_name)


def createPath(path):
    try:
        os.makedirs(path)
    except:
        pass

def mainTestCaseWise(arg1):
    (outputPdfPath,outputPdfName)=formatUserInputs(arg1)
    generateTestCaseWiseReport(outputPdfPath)


def writeTestCase(tcId,testDesc,status,startTime,endTime):
    #varTestCaseHeading
    #varTestCaseInfo
    #status - tc-pass

    systemConfig.testCaseWrittenOnce=True
    # with open('htmlTemplates/testCaseStart.html', 'r') as file:
    #     data = file.read()

    data="""<div class="accordion-item">
    <h2 class="accordion-header" id="{{varheadingOne}}">
        <button class="accordion-button collapsed {{status}}" type="button" data-bs-parent="#accordionExample" data-bs-toggle="collapse" data-bs-target="#{{varcollapseOne}}" aria-expanded="false" aria-controls="{{varcollapseOne}}">
            {{varTestCaseHeading}}
        </button>
    </h2>
    <div id="{{varcollapseOne}}" class="accordion-collapse collapse" aria-labelledby="{{varheadingOne}}" data-bs-parent="#accordionExample">
        <div class="accordion-body {{status}} test-scenario">
            <strong>{{varTestCaseInfo}}</strong>
        </div>

        <div>
            <div class="accordion test-step" id="{{varchildAccordionExample}}">
            </div>
            <div class="accordion-item test-step"></div>"""

    accordianId=generateRandomString()
    data=data.replace("{{varaccordionExample}}",accordianId)
    setChildAccordianId(accordianId)
    data=data.replace("{{varchildAccordionExample}}",getChildAccordianId())
    data=data.replace("{{varheadingOne}}",generateRandomString())
    data=data.replace("{{varcollapseOne}}",generateRandomString())
    data=data.replace("{{varTestCaseHeading}}","<xmp>{0} - {1}</xmp>".format(tcId,testDesc))
    data=data.replace("{{varTestCaseInfo}}","Start time : {0}<br>End time : {1}<br>".format(startTime, endTime))
    if not status:
        status="fail"
    if "pass" in status.lower():
        data=data.replace("{{status}}","tc-pass")
    else:
        data=data.replace("{{status}}","tc-fail")


    file=open("{0}".format(systemConfig.htmlFileName),"a")
    file.write(data)
    file.close()



def endTestCase():
    # with open('htmlTemplates/testCaseEnd.html', 'r') as file:
    #     data = file.read()

    data="""</div></div></div>"""

    file=open("{0}".format(systemConfig.htmlFileName),"a")
    file.write(data)
    file.close()


def generateRandomString():
    return "var"+''.join(random.choice(string.ascii_uppercase + string.digits) for _ in range(6))+str(int(time.time()))


def getChildAccordianId():
    return systemConfig.childAccordianId

def setChildAccordianId(accordianId):
    systemConfig.childAccordianId="child"+accordianId


def writeTestStep(testStepNumber,testDesc,expectedResult,actualResult,status):
    #varTestStepHeading
    #varTestStepInfo
    # with open('htmlTemplates/testStep.html', 'r') as file:
    #     data = file.read()

    data="""<h2 class="accordion-header test-step" id="{{varchildHeadingOne}}">
    <button class="accordion-button collapsed {{status}}" type="button" data-bs-parent="#{{varchildAccordionExample}}" data-bs-toggle="collapse" data-bs-target="#{{varchildCollapseOne}}" aria-expanded="false" aria-controls="{{varchildCollapseOne}}">
                            {{varTestStepHeading}}
                        </button>
</h2>
<div id="{{varchildCollapseOne}}" class="accordion-collapse collapse test-step" aria-labelledby="{{varchildHeadingOne}}" data-bs-parent="#{{varchildAccordionExample}}">
    <div class="accordion-body {{status}} limit-content">
        <strong>{{varTestStepInfo}}</strong>
    </div>
</div>"""

    data=data.replace("{{varTestStepHeading}}","<xmp>{0} - {1}</xmp>".format(testStepNumber,testDesc))
    data=data.replace("{{varTestStepInfo}}","<br><xmp>Expected Result : {0}</xmp><xmp> \nActual Result : {1}</xmp><br>".format(expectedResult,actualResult))

    data=data.replace("{{varchildAccordionExample}}",getChildAccordianId())
    data=data.replace("{{varchildHeadingOne}}",generateRandomString())

    randomChars=''.join(random.choice(string.ascii_uppercase + string.digits) for _ in range(6))
    data=data.replace("{{varchildCollapseOne}}",generateRandomString())


    if "pass" in status.lower():
        data=data.replace("{{status}}","tc-pass")
    else:
        data=data.replace("{{status}}","tc-fail")
    #data="<!--"+data+"-->"
    file=open("{0}".format(systemConfig.htmlFileName),"a")
    file.write(data)
    file.close()



def createHtmlHeader():
    print("Cur dir : {0}".format(os.getcwd()))
    # with open('htmlTemplates/header.html', 'r') as file:
    #     data = file.read()

    data="""<!doctype html>
<html lang="en">

<head>
    <!-- Required meta tags -->
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">

    <!-- Bootstrap CSS -->
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet" integrity="sha384-1BmE4kWBq78iYhFldvKuhfTAU6auU8tT94WrHftjDbrCEXSU1oBoqyl2QvZ6jIW3" crossorigin="anonymous">
    <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/js/bootstrap.bundle.min.js" integrity="sha384-ka7Sk0Gln4gmtz2MlQnikT1wXgYsOg+OMhuP+IlRH9sENBO0LRn5q+8nbTov4+1p" crossorigin="anonymous"></script>
    <!-- <link rel="stylesheet" href="styles.css"> -->
    <!-- <style type="text/css">
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        xmp {
            font-family: Trebuchet MS, Lucida Grande, Lucida Sans Unicode, Lucida Sans, Tahoma, sans-serif;
        }

        body {
            background-color: #212529;
            font-family: Trebuchet MS, Lucida Grande, Lucida Sans Unicode, Lucida Sans, Tahoma, sans-serif;
            min-height: 100vh;
            max-height: 100vh;
            max-width: 100vw;
            min-width: 100vw;
            margin: 0px;
            padding: 0px;
            float: left;
        }

        p {
            margin: 0;
            padding: 0;
        }

        .container,
        .container-lg,
        .container-md,
        .container-sm,
        .container-xl {
            min-height: 100vh;
            max-height: 100vh;
            max-width: 100vw;
            min-width: 100vw;
            margin: 0px;
            padding: 0px;
            background-color: #212529;
            float: left;
            position: absolute;
        }

        .test-step {
            margin-left: 50px;
        }

        .limit-content {
            min-width: 100vw;
            overflow: scroll;
        }

        .test-scenario {
            margin-left: 50px;
            margin-bottom: 10px;
        }

        .card {
            margin: 1rem;
        }

        .text-white {
            color: #fff;
        }

        .tc-pass {
            background-color: #bcf5bc;
        }

        .tc-fail {
            background-color: #ffb1b1;
        }

        .test-suite-info {
            background-color: #455168;
        }

        .greyish-text {
            color: #00000a;
        }

        .navbar-blue {
            background-color: #b1b1ff;
        }

        .custom-bold {
            font-weight: bold;
        }

        .custom-black-background {
            background-color: #000;
        }

        .custom-hyperlink-color {
            color: #000;
            font-weight: bold;
        }

        .custom-align-right {
            float: right;
        }

        .custom-card-text-background {
            background-color: #007bff;
        }

        .custom-align-center {
            float: center;
        }

        .custom-hyperlink-fontsize {
            font-size: 0.8rem;
            font-weight: bold;
        }

        body .custom-font-trebuchet {
            font-family: Trebuchet MS, Lucida Grande, Lucida Sans Unicode, Lucida Sans, Tahoma, sans-serif;
        }

        .center-on-screen {
            margin: auto;
            /* Added */
            /*float: none; /* Added */
            /* Added */
        }

        .big-font {
            font-size: 1.5rem;
        }

        .medium-font {
            font-size: 1rem;
        }

        .center-login {
            display: flex;
            justify-content: center;
            align-items: center;
            height: 200px;
            border: 3px solid green;
        }

        .flex-fill {
            flex: 1 1 auto;
        }

        .vertical-center {
            min-height: 100%;
            /* Fallback for browsers do NOT support vh unit */
            min-height: 100vh;
            /* These two lines are counted as one :-)       */
            display: flex;
            align-items: center;
        }

        .body-background {
            background-color: #020202;
        }

        .testCaseProps {
            border-top-width: 0px;
            border-bottom-width: 0px;
        }

        .accordion-button:not(.collapsed) {
            color: #FFF !important;
            background-color: #212529 !important;
        }

        .accordion-button:link,
        .accordion-button:visited,
        .accordion-button:hover,
        .accordion-button:active {
            background-color: #212529 !important;
            color: #FFF !important;
            text-decoration: none !important;
            border: hidden !important;
            border-color: #FFF !important;
            box-shadow: 0px !important;
        }

        .accordion-button:focus {
            z-index: 2;
            border-color: #FFF !important;
            outline: 0;
            /* box-shadow: 0 0 0 .25rem #212529 !important; */
        }
    </style> -->

    <style type="text/css">
        *:not(xmp, strong) {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
            background-color: #212529;
        }

        xmp {
            font-family: Trebuchet MS, Lucida Grande, Lucida Sans Unicode, Lucida Sans, Tahoma, sans-serif;
            font-weight: bold;
        }

        body {
            background-color: #212529;
            font-family: Trebuchet MS, Lucida Grande, Lucida Sans Unicode, Lucida Sans, Tahoma, sans-serif;
            min-height: 100vh;
            max-height: 100vh;
            max-width: 100vw;
            min-width: 100vw;
            margin: 0px;
            padding: 0px;
            float: center;
        }

        p {
            margin: 0;
            padding: 0;
        }

        .container,
        .container-lg,
        .container-md,
        .container-sm,
        .container-xl {
            min-height: 100vh;
            max-height: 100vh;
            max-width: 95vw;
            min-width: 95vw;
            margin: auto;
            padding: 0px;
        }

        .test-step {
            margin-left: 50px;
        }

        .limit-content {
            min-width: 100vw;
            overflow: scroll;
        }

        .test-scenario {
            margin-bottom: 10px;
        }

        .card {
            margin: 1rem;
        }

        .text-white {
            color: #fff;
        }

        .tc-pass {
            background-color: #bcf5bc;
        }

        .tc-fail {
            background-color: #ffb1b1;
        }

        .test-suite-info {
            background-color: #455168;
        }

        .greyish-text {
            color: #00000a;
        }

        .navbar-blue {
            background-color: #b1b1ff;
        }

        .custom-bold {
            font-weight: bold;
        }

        .custom-black-background {
            background-color: #000;
        }

        .custom-hyperlink-color {
            color: #000;
            font-weight: bold;
        }

        .custom-align-right {
            float: right;
        }

        .custom-card-text-background {
            background-color: #007bff;
        }

        .custom-align-center {
            float: center;
        }

        .custom-hyperlink-fontsize {
            font-size: 0.8rem;
            font-weight: bold;
        }

        body .custom-font-trebuchet {
            font-family: Trebuchet MS, Lucida Grande, Lucida Sans Unicode, Lucida Sans, Tahoma, sans-serif;
        }

        .center-on-screen {
            margin: auto;
            /* Added */
            /*float: none; /* Added */
            /* Added */
        }

        .big-font {
            font-size: 1.5rem;
        }

        .medium-font {
            font-size: 1rem;
        }

        .center-login {
            display: flex;
            justify-content: center;
            align-items: center;
            height: 200px;
            border: 3px solid green;
        }

        .flex-fill {
            flex: 1 1 auto;
        }

        .vertical-center {
            min-height: 100%;
            /* Fallback for browsers do NOT support vh unit */
            min-height: 100vh;
            /* These two lines are counted as one :-)       */
            display: flex;
            align-items: center;
        }

        .body-background {
            background-color: #020202;
        }

        .testCaseProps {
            border-top-width: 0px;
            border-bottom-width: 0px;
        }

        .accordion-button:not(.collapsed) {
            color: #FFF !important;
            background-color: #212529 !important;
        }

        .accordion-button:link,
        .accordion-button:visited,
        .accordion-button:hover,
        .accordion-button:active {
            background-color: #212529 !important;
            color: #FFF !important;
            text-decoration: none !important;
            border: hidden !important;
            /* border-color: #FFF !important; */
            box-shadow: 0px !important;
            border: 0px;
        }

        .accordion-button:focus {
            z-index: 2;
            border-color: #FFF !important;
            outline: 0;
            /* box-shadow: 0 0 0 .25rem #212529 !important; */
        }
    </style>


    <title>Execution Report</title>
</head>

<body>
    <div class="container bg-dark">

        <h1 style="text-align:center;"><img style=" text-align: center;margin-right: 20px;" src="https://drive.google.com/thumbnail?id=1qxlriskSJV_pIgZDjU-Vxy5Jb0CaIxdX " alt="" width="150 " height="80" /><img style=" text-align: center;margin-left: 20px;" src="https://drive.google.com/thumbnail?id=1GYdEyzQ7OaxGpe9LCvpx_lPFdONJ0L9o"
                alt=" " width="150 " height="80" /></h1>
        <h4 style="text-align:center;margin-top: 20px;" class="text-white">Automated Execution Summary</h1>
            <h5 style="text-align:center;margin-top: 20px;margin-bottom: 30px;" class="text-white">Total TCs : 10 | Passed : 10 | Failed : 0</h1>

                <div class="accordion" id="accordionExample">"""

    file=open("{0}".format(systemConfig.htmlFileName),"w")
    file.write(data)
    file.close()



def createHtmlFooter():
    # with open('htmlTemplates/footer.html', 'r') as file:
    #     data = file.read()

    data="""</div>
</div>
</div>
</div>

</body>

</html>"""

    file=open("{0}".format(systemConfig.htmlFileName),"a")
    file.write(data)
    file.close()


def writeTestCaseCountInHtmlReport():
    filename=systemConfig.htmlFileName
    text_to_search='Total TCs : 10 | Passed : 10 | Failed : 0'
    replacement_text='Total TCs : {0} | Passed : {1} | Failed : {2}'.format(systemConfig.pass_ctr+systemConfig.fail_ctr,systemConfig.pass_ctr,systemConfig.fail_ctr)
    with open(filename, 'r') as file:
            data = file.read()

    data=data.replace(text_to_search,replacement_text)
    file=open(filename,"w")
    file.write(data)
    file.close()


def mainEntireReport(absExcelName):
    (outputPdfPath,outputPdfName)=formatUserInputs(absExcelName)
    dataRangeToConvert=getExcelRangeWithoutHeaders()
    initializeWordReport()
    loopCounter=0
    testCaseLogged=False

    createHtmlHeader()


    for eachExcelRow in dataRangeToConvert:
        #print("Excel Row is : {0}".format(eachExcelRow))
        if rowIsTestCase(eachExcelRow):
            createTestCaseHeaderAndPopulate(eachExcelRow)
            testCaseLogged=True

        else:
            if testCaseLogged is True:
                testCaseLogged=False
                createTestStepHeaderAndPopulate(eachExcelRow)
            else:
                if systemConfig.testCaseWrittenOnce:
                    populateTestStepDetails(eachExcelRow)
            addScreenshots(eachExcelRow)

    formatParagraphs()
    endWordReporting()


    createHtmlFooter()
    writeTestCaseCountInHtmlReport()
    #saveEntireWordDocument(outputPdfPath,outputPdfName)



def main(arg1):
    #arg1 is the abs file name

    userConfig.excelFileName=arg1
    userConfig.excelFilePath=(arg1.split("\\")[:-1]) #['C:', 'Users', 'ALTAMASH', 'Desktop', 'TCoE', 'API', '1Gie', 'Results', 'Result_2021_11_07_01_32_43715000']
    userConfig.excelFilePath=os.path.join(*userConfig.excelFilePath).replace(":",":/")
    systemConfig.htmlFileName=userConfig.excelFileName.replace(".xlsx",".html")
    #systemConfig.htmlFileName="{0}/Execution_Report_{1}.html".format(userConfig.excelFilePath,time.time())
    arg6="pc"
    if "pc" in arg6.lower():
            Configuration_Browser.deviceIsMobile=False
    else:
        Configuration_Browser.deviceIsMobile=True

    mainEntireReport(arg1)
    #mainTestCaseWise(arg1,arg2,arg3,arg4,arg5,arg6,arg7)





