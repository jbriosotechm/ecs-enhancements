from docx import Document
from docx.shared import Inches,Pt,Length
import openpyxl
from openpyxl import styles
from openpyxl.utils import get_column_letter,column_index_from_string
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import time
import os,sys
import datetime
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil
import re
from Configuration_Browser import *


import shutil

wdFormatPDF = 17
pass_ctr=0
fail_ctr=0
hexcolorcode_first_row=hexcolorcode_test_case_ref
columns_to_ignore=['B','D','E','F','G']






################################################################################
################################################################################
################################################################################


def GetCellColor_based_on_color(cellObj):
    global sheet
    global hexcolorcode_fail
    global hexcolorcode_pass
    global pass_ctr
    global fail_ctr
    Colors=styles.colors.COLOR_INDEX
    cell_color=sheet[cellObj.coordinate].fill.start_color.index

    if len(str(cell_color))>=6:
        return str(cell_color)[2:]
    else:
        return (Colors[cell_color])[2:]


def GetCellColor(cellObj):
    global sheet
    global hexcolorcode_fail
    global hexcolorcode_pass
    global pass_ctr
    global fail_ctr


    row_number=str(re.sub("[^0-9]", "", cellObj.coordinate))

    if str(row_number)=='1':
        return '000000'

    StatusCell=StatusColumn+row_number
    result_test_case=sheet[StatusCell].value

    try:
        result_test_case=result_test_case.lower()
    except:
        try:
            if (sheet['A'+row_number].value).strip():
                fail_ctr+=1
        except:
            pass

        return hexcolorcode_fail

    if 'pass' in result_test_case:
        try:
            if (sheet['A'+row_number].value).strip():
                pass_ctr+=1
        except:
            pass

        return hexcolorcode_pass

    try:
        if (sheet['A'+row_number].value).strip():
            fail_ctr+=1
    except:
        pass
    return hexcolorcode_fail





def ColorTheCell(table,word_column_number,hexcolorcode):

    data=r'<w:shd {} w:fill="'+hexcolorcode+r'"/>'
    table.rows[0].cells[word_column_number]._tc.get_or_add_tcPr().append(parse_xml(data.format(nsdecls('w'))))


##############################################################
##############################################################
exception_para=[]


def main(arg1,arg2,arg3,arg4,arg5):
    global sheet
    global hexcolorcode_fail
    global hexcolorcode_pass
    global pass_ctr
    global fail_ctr
    ###############################################################
    ###############################################################
    # Parameters needed for running
    # Name of input excel
    # path of input excel
    # name of output PDF
    # path of output PDF
    # path of python

    ###############################################################

    script_path = arg1
    input_excel_name = arg2
    input_excel_path = arg3
    output_pdf_path = arg4
    output_pdf_name = arg5.split('.')[0]

    print "Script path :", script_path
    print "input_excel_name: ", input_excel_name
    print "input_excel_path", input_excel_name
    print "output_pdf_path", output_pdf_path
    print "output_pdf_name", output_pdf_name


    index = 1


    if script_path[-1] == "\\" and script_path[-2] == "\\":
        script_path = script_path[:-1]
    print "Script path : ", script_path

    if output_pdf_path[-1] == "\\" and output_pdf_path[-2] == "\\":
        output_pdf_path = output_pdf_path[:-1]
    print "output_pdf_path : ", output_pdf_path

    document = Document()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n")
    run.add_picture(globe_logo, Inches(2), Inches(1))
    run.add_tab()
    run.add_picture(techm_logo, Inches(2), Inches(1))

    heading_para = document.paragraphs[-1]
    paragraph = document.add_paragraph('\nTest Summary Report')

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # paragraph.add_run('Test Summary Report', style = 'Title')
    paragraph.style = document.styles['Title']

    stats_paragraph = document.add_paragraph('ReplaceContent')
    stats_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stats_paragraph.style = document.styles['Heading 1']

    section = document.sections[-1]
    sections = document.sections
    section.page_width = Inches(11)
    section.page_height = Inches(17)

    done_once = False

    # margins will also be a config parameter
    for section in sections:
        section.top_margin = Inches(0.05)
        section.bottom_margin = Inches(0.2)
        section.left_margin = Inches(0.2)
        section.right_margin = Inches(0.2)
        section.orientation = WD_ORIENT.LANDSCAPE

    ##############################################################
    ##############################################################
    ##############################################################

    wb = openpyxl.load_workbook(input_excel_name)
    sheet = wb.active
    excel_number_of_rows = sheet.max_row

    excel_number_of_columns = sheet.max_column
    excel_last_column_name = get_column_letter(excel_number_of_columns)
    screenshot_column_name = 'I'
    # screenshot_column_number=column_index_from_string('I')

    current_row = 0
    screenshot_path = None

    ##############################################################
    ##############################################################
    ##############################################################

    #############################################
    # iterate over all rows
    #############################################
    print "Will enter while loop"
    while current_row < excel_number_of_rows:

        cell_color = None
        current_row += 1

        print "Evaluating row : ", current_row

        from_cell = "A" + str(current_row)
        to_cell = excel_last_column_name + str(current_row)
        whole_row_data = sheet[from_cell:to_cell]
        # till here the current row data has been extracted

        screenshot_cell = screenshot_column_name + str(current_row)

        #############################################
        # handling for all rows except first row
        #############################################

        if current_row > 1:
            # evaluating if SS is available
            try:
                screenshot_path = (sheet[screenshot_cell].value).strip()
            except:
                screenshot_path = None

            word_column_number = 0
            word_column_number1 = 0

            first_cell_val = None
            first_cell_name = None

            for rowOfCellObjects in whole_row_data:

                first_cell_name = (rowOfCellObjects[0]).coordinate

                try:
                    # handling for test case - if first cell of row is not empty, it
                    # implies that it is a test case

                    first_cell_val = ((rowOfCellObjects[0]).value).strip()

                    # add table header also
                    document.add_paragraph('\n')
                    table1 = document.add_table(rows=1, cols=excel_number_of_columns - reduce_columns - 3)
                    table1.style = 'TableGrid'
                    table1.allow_autofit = True
                    hdr_cells1 = table1.rows[0].cells

                    temp_excel_column_indicator = 1

                    for cellObj in rowOfCellObjects:
                        if re.sub("[0-9]", "", cellObj.coordinate) not in [screenshot_column_name,
                                                                           Test_Object_Name_Column, Test_Data_Column,
                                                                           'B', 'F', 'G']:
                            hdr_cells1[word_column_number1].paragraphs[0].add_run(
                                (sheet[get_column_letter(temp_excel_column_indicator) + "1"]).value).bold = True
                            ColorTheCell(table1, word_column_number1, hexcolorcode_first_row)
                            word_column_number1 += 1

                        temp_excel_column_indicator += 1

                    word_column_number = 0
                    word_column_number1 = 0
                    table3 = document.add_table(rows=1, cols=excel_number_of_columns - reduce_columns - 3)
                    table3.style = 'TableGrid'
                    table3.allow_autofit = True
                    hdr_cells3 = table3.rows[0].cells

                    # document.add_paragraph('hello world')
                except:
                    # handling for test step

                    if not done_once:
                        done_once = True
                        # document.add_paragraph('hello again')

                        temp_excel_column_indicator = 0
                        table2 = document.add_table(rows=1, cols=len(test_step_reference))
                        table2.style = 'TableGrid'
                        table2.allow_autofit = True
                        hdr_cells2 = table2.rows[0].cells
                        for step_ref in test_step_reference:
                            # print "Step Ref is : ",step_ref
                            hdr_cells2[temp_excel_column_indicator].paragraphs[0].add_run(step_ref).bold = True
                            ColorTheCell(table2, temp_excel_column_indicator, hexcolorcode_test_step_ref)
                            temp_excel_column_indicator += 1

                        temp_excel_column_indicator = 0

                    first_cell_val = None
                    table4 = document.add_table(rows=1, cols=excel_number_of_columns - reduce_columns - 3)
                    table4.style = 'TableGrid'
                    table4.allow_autofit = True
                    hdr_cells4 = table4.rows[0].cells

                hexcolorcode = GetCellColor(rowOfCellObjects[0])
                # print "Hex color code is : ",hexcolorcode

                is_main_step = False
                main_step = ""
                for cellObj in rowOfCellObjects:
                    # process all columns except the one that has SS path
                    if screenshot_column_name not in cellObj.coordinate and Test_Object_Name_Column not in cellObj.coordinate and Test_Data_Column not in cellObj.coordinate:
                        try:
                            if first_cell_val is None:
                                # handling for Test Steps
                                if not is_main_step:
                                    hdr_cells4[word_column_number].text = cellObj.value

                                if 'C' in cellObj.coordinate:
                                    if '[Main Step]' in cellObj.value:
                                        main_step = cellObj.value.replace("[Main Step]", "")
                                        is_main_step = True

                                if 'F' in cellObj.coordinate:
                                    main_test_step_description = ""
                                    if cellObj.value is not None:
                                        main_test_step_description = cellObj.value

                                if 'H' in cellObj.coordinate:
                                    if is_main_step:
                                        hdr_cells4[1].text = main_step
                                        hdr_cells4[2].text = main_test_step_description
                                        hdr_cells4[4].text = cellObj.value
                                        if 'fail' in cellObj.value.lower():
                                            color = "faa0a0"
                                        else:
                                            color = "c1e1c1"

                                        ColorTheCell(table4,0,color)
                                        ColorTheCell(table4,1,color)
                                        ColorTheCell(table4,2,color)
                                        ColorTheCell(table4,3,color)
                                        ColorTheCell(table4,4,color)
                            else:
                                # handling for Test Cases

                                if 'B' not in cellObj.coordinate and 'F' not in cellObj.coordinate and 'G' not in cellObj.coordinate:
                                    done_once = False
                                    hdr_cells3[word_column_number1].paragraphs[0].add_run(cellObj.value).bold = True
                                    # print "Cell obj is : ",cellObj.coordinate
                                    ColorTheCell(table3, word_column_number1, hexcolorcode)
                                    word_column_number1 += 1



                        except:
                            try:
                                word_column_number -= 1
                            except:
                                pass

                        word_column_number += 1

        if screenshot_path is not None:
            try:
                if auto_image_height:
                    document.add_paragraph('\n')
                    document.add_picture(screenshot_path, width=Inches(2.5))
                    last_paragraph = document.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    document.add_paragraph('\n')
                else:
                    document.add_paragraph('\n')
                    try:
                        if (screenshot_path.split("\\")[-1]).lower().startswith("mobile_"):
                            print "Mobile SS scenario"
                            document.add_picture(screenshot_path, width=Inches(image_width_inches_mobile),
                                                 height=Inches(image_height_inches_mobile))
                        else:
                            document.add_picture(screenshot_path, width=Inches(image_width_inches),
                                                 height=Inches(image_height_inches))
                    except Exception, e:
                        print "Exception : ", e
                        print "Setting default to browser dimensions"
                        document.add_picture(screenshot_path, width=Inches(image_width_inches),
                                             height=Inches(image_height_inches))

                    last_paragraph = document.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    document.add_paragraph('\n')
            except Exception, e:
                pass
                # document.add_paragraph('Unable to add screenshot, Exception is : '+str(e))
                try:
                    pass
                    # exception_para.append(document.paragraphs[-1])
                    # exception_para.append(document.paragraphs[-2])
                except:
                    pass
                # time.sleep(5)

        para_counter = 0
        for para in document.paragraphs:
            if para_counter > 1:
                paragraph_format = para.paragraph_format
                paragraph_format.line_spacing = 0.3
            else:
                paragraph_format = para.paragraph_format
                paragraph_format.line_spacing = 0.7
            para_counter += 1

        # paragraph_format=heading_para.paragraph_format
        # paragraph_format.line_spacing = 7

        if len(exception_para) > 0:
            for para in exception_para:
                paragraph_format = para.paragraph_format
                paragraph_format.line_spacing = 1.25

    # add_run(text=None, style=None)[source]

    total_test_cases = pass_ctr + fail_ctr
    result_str = "Total Test Cases : {0}      Passed : {1}      Failed : {2}".format(total_test_cases, pass_ctr,
                                                                                     fail_ctr)
    # stats_paragraph.add_run(result_str,style='Subtitle')
    stats_paragraph.text = result_str
    paragraph_format = stats_paragraph.paragraph_format
    paragraph_format.line_spacing = 2
    current_time = str(datetime.datetime.now()).replace(' ', '_').replace('-', '_').replace(':', '')
    document.save(output_pdf_path.replace(".pdf",".docx"))
    print "Word Report Generated"

